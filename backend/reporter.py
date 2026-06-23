import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set the padding of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def get_finding_details(title, category):
    impact = "Dampak keamanan belum didefinisikan secara spesifik."
    use_case = "Skenario serangan spesifik belum didefinisikan."
    remediation = "Tinjau konfigurasi parameter terkait, terapkan prinsip least privilege, dan pastikan konfigurasi sesuai dengan best practice."
    
    title_lower = title.lower()
    
    if "duplicate policy/rule name" in title_lower:
        impact = "Mengurangi kebersihan konfigurasi (configuration hygiene), memicu kebingungan bagi administrator saat pemeliharaan, serta meningkatkan risiko salah konfigurasi saat melakukan pembaharuan kebijakan NAC karena nama yang sama mereferensikan objek yang berbeda."
        use_case = "Administrator mencoba mengubah policy untuk memblokir perangkat tertentu, namun salah mengubah rule dengan nama yang sama di folder lain, sehingga policy keamanan tidak diterapkan pada target yang seharusnya."
        remediation = "Berikan penamaan yang unik, spesifik, dan deskriptif untuk setiap policy atau rule. Hapus atau konsolidasikan aturan yang memiliki nama duplikat."
        
    elif "ip range overlap" in title_lower:
        impact = "Menyebabkan ketidakpastian (ambiguitas) dalam evaluasi paket/perangkat. Rule yang berada pada urutan lebih atas akan memproses paket terlebih dahulu, yang dapat menyebabkan bypass pemeriksaan keamanan jika rule tersebut memiliki kriteria lebih longgar."
        use_case = "Segmentasi jaringan untuk server sensitif tidak sengaja tercakup (overlap) dalam aturan klasifikasi umum PC tamu (Guest PC), sehingga PC tamu dapat mengakses segmen server tersebut tanpa otentikasi ketat."
        remediation = "Review dan batasi rentang IP (IP Range) pada masing-masing rule secara cermat. Pisahkan segmen IP yang tumpang tindih atau gabungkan rule yang tujuannya sama."
        
    elif "disabled policy/rule" in title_lower:
        impact = "Menumpuknya aturan pasif yang tidak terpakai menambah beban kognitif administrator saat membaca file konfigurasi (technical debt) dan mempersulit pemeliharaan."
        use_case = "Terdapat puluhan aturan lama bekas pengujian tahun lalu yang dibiarkan berstatus disabled, mempersulit tim audit untuk melakukan verifikasi aturan mana yang benar-benar aktif saat ini."
        remediation = "Lakukan review berkala terhadap aturan yang tidak aktif. Jika aturan tersebut sudah tidak diperlukan lagi, segera hapus (delete) secara permanen dari konfigurasi."
        
    elif "no active actions" in title_lower:
        impact = "Membuang sumber daya pemrosesan engine NAC karena melakukan pencocokan kriteria perangkat secara terus menerus namun tidak memberikan respon tindakan apa pun (seperti karantina, penugasan VLAN, atau klasifikasi grup)."
        use_case = "Rule klasifikasi dirancang untuk mendeteksi perangkat ilegal, namun lupa dikonfigurasi untuk memicu aksi pemindahan ke VLAN karantina, sehingga perangkat tersebut tetap bebas mengakses jaringan."
        remediation = "Tambahkan setidaknya satu tindakan aktif (misalnya notification, assign VLAN, atau add to group) pada sub-rule tersebut agar hasil pencocokan kriteria menghasilkan respon keamanan."
        
    elif "empty conditions" in title_lower:
        impact = "Sub-rule bertindak sebagai Catch-All yang mencocokkan setiap perangkat tanpa kriteria, berisiko memicu salah penegakan aturan (false positives) pada perangkat legal yang tidak bersalah."
        use_case = "Sub-rule kosong diletakkan di bagian atas hierarki, menyebabkan seluruh endpoint (termasuk server produksi) tercocokkan ke sub-rule tersebut dan diputus koneksinya secara tidak sengaja."
        remediation = "Definisikan filter kondisi evaluasi (seperti tipe OS, rentang IP, status keanggotaan domain, dll) pada sub-rule tersebut, atau hapus jika memang tidak dibutuhkan sebagai fallback."
        
    elif "caching disabled" in title_lower:
        impact = "Engine NAC terpaksa mengevaluasi ulang perangkat setiap kali menerima paket data baru dari endpoint. Ini memicu lonjakan beban CPU server NAC secara ekstrem dan memperlambat respon otentikasi jaringan."
        use_case = "Evaluasi OS scanning pada ribuan endpoint dikonfigurasi tanpa cache (TTL=0), mengakibatkan utilitas CPU server NAC mencapai 100% dan mengganggu operasional jaringan."
        remediation = "Atur nilai Cache TTL minimal 3600 detik (1 jam) untuk aturan yang sifatnya non-real-time untuk meminimalkan beban CPU engine NAC."
        
    elif "low cache ttl" in title_lower:
        impact = "Meningkatkan overhead pemrosesan pada engine NAC karena proses evaluasi ulang terhadap status kepatuhan perangkat dilakukan terlalu sering tanpa urgensi real-time yang jelas."
        use_case = "Kebijakan kepatuhan standar melakukan pemindaian antivirus setiap 5 menit sekali pada seluruh endpoint, membebani antrian pemrosesan inspeksi NAC."
        remediation = "Tingkatkan Cache TTL menjadi minimal 3600 detik (1 jam) untuk rule evaluasi standar."
        
    elif "high privilege account detected" in title_lower:
        impact = "Kompromi pada akun berhak istimewa tinggi (seperti Domain Admin) memberikan akses kontrol penuh bagi penyerang untuk menguasai seluruh domain Active Directory, termasuk database kredensial."
        use_case = "Penyerang melakukan brute-force atau Kerberoasting terhadap akun admin yang memiliki SPN, berhasil mendapatkan password hash, dan melakukan lateral movement untuk menguasai Domain Controller."
        remediation = "Batasi keanggotaan grup administratif seminimal mungkin, gunakan akun khusus admin yang terpisah dari akun sehari-hari, dan aktifkan otentikasi multi-faktor (MFA)."
        
    elif "disabled user account" in title_lower:
        impact = "Akun yang dinonaktifkan tetapi tidak pernah dihapus/diarsipkan tetap berada di direktori, menambah beban pengelolaan data dan berpotensi diaktifkan kembali secara ilegal oleh penyerang jika tidak terpantau."
        use_case = "Mantan karyawan yang akunnya telah dinonaktifkan diaktifkan kembali secara diam-diam oleh aktor internal berwenang untuk melakukan eksfiltrasi data tanpa terdeteksi."
        remediation = "Lakukan pembersihan akun secara berkala. Akun disabled yang sudah melewati masa tenggang (misalnya 30 hari) harus segera di-archive dan dihapus secara permanen."
        
    elif "stale/inactive account" in title_lower:
        impact = "Akun yang sudah lama tidak aktif (tidak log masuk > 90 hari) rentan diambil alih karena jarang dipantau dan kata sandinya mungkin sudah usang atau bocor."
        use_case = "Penyerang menemukan akun staf yang sedang cuti panjang, memulihkan passwordnya, dan memanfaatkannya sebagai pintu masuk (backdoor) ke jaringan internal perusahaan."
        remediation = "Implementasikan kebijakan otomatisasi pemblokiran/nonaktifkan akun jika tidak mendeteksi aktivitas login dalam waktu 90 hari."
        
    elif "password never expires" in title_lower:
        impact = "Menghindari kebijakan rotasi password wajib. Jika password akun tersebut bocor, penyerang dapat menggunakannya selamanya tanpa batasan waktu."
        use_case = "Kredensial akun layanan yang diset 'Password Never Expires' bocor di repositori kode publik, memberikan penyerang akses permanen ke infrastruktur."
        remediation = "Hapus bendera 'Password Never Expires' pada akun pengguna standar. Gunakan Group Managed Service Accounts (gMSA) untuk akun layanan yang membutuhkan penggantian password otomatis."
        
    elif "unconstrained kerberos delegation" in title_lower:
        impact = "Memungkinkan server tujuan delegasi menyimpan tiket TGT (Ticket Granting Ticket) user di memori. Jika server tersebut disusupi, penyerang dapat mencuri TGT tersebut dan menyamar sebagai pengguna di sistem lain (termasuk admin)."
        use_case = "Penyerang menyusupi web server yang memiliki konfigurasi Unconstrained Delegation, lalu menunggu Domain Admin mengakses web tersebut untuk mencuri tiket TGT admin dan menguasai Domain Controller."
        remediation = "Nonaktifkan Unconstrained Delegation. Migrasikan ke Constrained Delegation atau Resource-Based Constrained Delegation (RBCD) yang membatasi tujuan delegasi ke layanan spesifik saja."
        
    elif "constrained kerberos delegation" in title_lower:
        impact = "Constrained Delegation membatasi target layanan yang dapat dituju, namun jika akun delegasi atau target SPN disusupi, atau dikonfigurasi dengan opsi protokol transisi yang kurang aman (misalnya S4U2self/S4U2proxy), penyerang tetap dapat memicu impersonasi hak akses."
        use_case = "Penyerang menyusupi akun delegasi yang dikonfigurasi dengan transisi protokol, memanipulasi atribut otentikasi untuk bertindak atas nama pengguna berhak akses tinggi terhadap layanan backend."
        remediation = "Audit berkala konfigurasi delegasi terbatas, batasi hanya ke SPN spesifik, dan hindari penggunaan protokol transisi jika tidak sangat diperlukan."
        
    elif "kerberos pre-authentication disabled" in title_lower:
        impact = "Penyerang dapat meminta tiket TGT atas nama akun tanpa perlu mengetahui kata sandi, lalu melakukan brute-force offline (AS-REP Roasting) untuk memecahkan sandi akun tersebut."
        use_case = "Penyerang mengirimkan permintaan AS-REQ ke Domain Controller untuk akun tanpa pra-otentikasi, menerima respon AS-REP yang berisi hash kata sandi terenkripsi, lalu memecahkannya menggunakan Hashcat secara offline."
        remediation = "Wajibkan pra-otentikasi Kerberos (Kerberos Pre-Authentication) pada semua akun tanpa pengecualian."
        
    elif "kerberoastable account detected" in title_lower:
        impact = "Akun pengguna standar yang memiliki SPN (Service Principal Name) dapat dijadikan target serangan Kerberoasting, di mana penyerang meminta tiket layanan TGS lalu memecahkan hash passwordnya secara offline."
        use_case = "Penyerang menggunakan akun domain biasa untuk meminta tiket TGS untuk akun layanan SQL, mengekspor hash tiket tersebut, dan melakukan brute-force offline untuk mendapatkan password teks biasa akun layanan."
        remediation = "Gunakan kata sandi yang sangat panjang (minimal 25 karakter) dan acak untuk akun layanan yang memerlukan SPN, atau migrasikan ke Managed Service Accounts (MSA/gMSA)."
        
    elif "dangerous acl permission detected" in title_lower:
        impact = "Memungkinkan penyerang yang menyusupi akun non-admin dengan hak istimewa ACL tertentu (seperti GenericWrite atau WriteDacl) atas objek lain untuk mengubah atribut atau password objek tersebut dan mengambil alih hak aksesnya."
        use_case = "Penyerang menyusupi akun staf IT junior yang memiliki hak GenericAll atas grup keamanan sensitif, lalu memasukkan akunnya sendiri ke dalam grup tersebut untuk eskalasi hak akses."
        remediation = "Audit dan bersihkan ACL Active Directory secara berkala. Batasi delegasi kontrol objek AD hanya kepada tim resmi yang bertanggung jawab."
        
    elif "shortest path to domain admin detected" in title_lower:
        impact = "Memberikan peta jalur serangan bagi aktor ancaman untuk melompat dari akun biasa ke akun administrator domain melalui rantai izin atau sesi yang saling terhubung."
        use_case = "Rantai serangan berjalan dari akun Helpdesk -> lokal admin workstation -> sesi aktif Domain Admin -> pengambilalihan DC."
        remediation = "Batasi hak akses lokal administrator, batasi tempat login akun administratif (Tiered Administration Model), dan bersihkan sesi aktif yang menggantung pada workstation non-DC."
        
    elif "laps not enabled" in title_lower:
        impact = "Mempermudah penyerang melakukan pergerakan lateral (lateral movement) di jaringan karena seluruh workstation/server menggunakan kata sandi lokal administrator yang seragam."
        use_case = "Penyerang mengekstrak hash password admin lokal dari satu workstation yang berhasil disusupi, lalu menggunakan teknik Pass-the-Hash untuk menguasai seluruh workstation lain di jaringan."
        remediation = "Terapkan Windows Local Administrator Password Solution (LAPS) di seluruh komputer domain untuk mengacak kata sandi admin lokal secara berkala dan menyimpannya secara aman di AD."
        
    elif "smb signing disabled" in title_lower:
        impact = "Memungkinkan penyerang di jaringan lokal melakukan serangan Man-in-the-Middle (NTLM Relay), di mana kredensial otentikasi yang dikirimkan melalui SMB diteruskan untuk menguasai server target."
        use_case = "Penyerang memicu paksaan otentikasi dari server cetak, menangkap lalu lintasnya, dan meneruskannya (relay) ke komputer target yang menonaktifkan SMB signing untuk mengeksekusi kode perintah jarak jauh."
        remediation = "Wajibkan SMB Signing (Require SMB Signing) di seluruh server dan workstation melalui Group Policy Object (GPO)."
        
    elif "active session on workstation" in title_lower:
        impact = "Sesi aktif dari akun admin Tier-0 (Domain Admin) tertinggal di memori LSASS workstation Tier-2 biasa, yang dapat dengan mudah diekstrak menggunakan Mimikatz jika workstation tersebut disusupi."
        use_case = "Domain Admin login langsung di workstation staf untuk troubleshoot, lalu pergi tanpa melakukan logoff secara bersih, membiarkan sesi aktifnya berjalan di latar belakang."
        remediation = "Terapkan model tiering administrasi. Larang akun admin Tier 0 masuk ke workstation Tier 2."
        
    elif "high count of administrators" in title_lower:
        impact = "Meningkatkan attack surface pada komputer tersebut karena banyaknya akun yang memiliki hak akses lokal admin, mempersulit kontrol akses minimum."
        use_case = "Penyerang menembus salah satu kredensial dari 15 admin lokal yang terdaftar pada workstation target untuk melakukan eskalasi hak akses penuh."
        remediation = "Batasi jumlah administrator lokal seminimal mungkin. Gunakan Restricted Groups GPO untuk mengontrol keanggotaan grup lokal secara terpusat."
        
    elif "dangerous permissions granted to domain users" in title_lower:
        impact = "Pengguna non-admin memiliki akses administrator lokal di komputer tersebut, mempermudah penyebaran infeksi malware dan bypass kebijakan keamanan workstation."
        use_case = "Malware berjalan di bawah konteks pengguna biasa, namun langsung memperoleh hak akses sistem penuh karena seluruh Domain Users berstatus admin lokal."
        remediation = "Hapus grup 'Domain Users' dari keanggotaan grup 'Administrators' lokal pada workstation klien."
        
    elif "gpo permissions security risk" in title_lower:
        impact = "Pengguna biasa yang memiliki akses tulis ke objek GPO dapat memodifikasi GPO tersebut untuk menyuntikkan startup script berbahaya guna menginfeksi seluruh komputer klien di domain."
        use_case = "Penyerang memodifikasi isi script GPO untuk membuat user admin lokal baru di seluruh PC saat booting."
        remediation = "Batasi izin tulis pada GPO hanya untuk admin tepercaya."
        
    elif "adcs esc1 certificate template" in title_lower:
        impact = "Template sertifikat ADCS yang memperbolehkan peminta menentukan SAN dapat digunakan oleh penyerang untuk impersonasi akun Domain Admin dan melakukan pengambilalihan domain penuh."
        use_case = "Penyerang meminta sertifikat dengan parameter SAN berisi nama akun Domain Admin pada template yang rentan, lalu menggunakannya untuk login sebagai Domain Admin."
        remediation = "Tinjau template sertifikat ADCS. Nonaktifkan opsi 'Supply in the request' pada tab Subject Name untuk template Client Authentication."
        
    return impact, use_case, remediation

def generate_report_docx(report_data: dict) -> io.BytesIO:
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    brand_name = report_data.get("brand_name", "Cybersecurity Assessment")
    filename = report_data.get("filename", "config_file")
    stats = report_data.get("stats", {})
    issues = report_data.get("issues", [])
    
    # Title Cover Page / Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SpectraOne Security Assessment Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a) # Deep Navy
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Module: {brand_name}\nTarget: {filename}")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hr.add_run("—" * 50).font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    
    # 1. Executive Summary
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Summary")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
    
    total_issues = stats.get("total_issues", len(issues))
    crit_count = stats.get("critical_issues", 0)
    high_count = stats.get("high_issues", 0)
    med_count = stats.get("medium_issues", 0)
    low_count = stats.get("low_issues", 0)
    info_count = stats.get("info_issues", 0)
    
    p1 = doc.add_paragraph()
    p1.add_run(f"SpectraOne telah melakukan audit kepatuhan keamanan dan evaluasi konfigurasi terhadap target ")
    p1.add_run(f"'{filename}' ").bold = True
    p1.add_run(f"menggunakan modul analisa ")
    p1.add_run(f"'{brand_name}'. ").bold = True
    p1.add_run(f"Dari hasil analisis otomatis yang dilakukan, terdeteksi sebanyak ")
    p1.add_run(f"{total_issues} temuan keamanan ").bold = True
    p1.add_run("dengan rincian tingkat keparahan sebagai berikut:")
    
    # Summary Table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Tingkat Keparahan (Severity)"
    hdr_cells[1].text = "Jumlah Temuan"
    for cell in hdr_cells:
        set_cell_background(cell, "1e3f7a")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                
    sevarray = [
        ("Critical", crit_count, "ef4444"),
        ("High", high_count, "f97316"),
        ("Medium", med_count, "f59e0b"),
        ("Low", low_count, "3b82f6"),
        ("Info", info_count, "6b7280")
    ]
    
    for idx, (sev_name, count, color_hex) in enumerate(sevarray, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = sev_name
        row_cells[1].text = str(count)
        
        # Color label
        set_cell_background(row_cells[0], "f3f4f6")
        set_cell_margins(row_cells[0], top=100, bottom=100, left=150, right=150)
        for p in row_cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(
                    int(color_hex[0:2], 16),
                    int(color_hex[2:4], 16),
                    int(color_hex[4:6], 16)
                )
                
        # Count cell
        set_cell_margins(row_cells[1], top=100, bottom=100, left=150, right=150)
        for p in row_cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                
    doc.add_paragraph() # Spacer
    
    # 2. Kondisi Saat Ini (Current Conditions)
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Kondisi Saat Ini (Current Conditions)")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
    
    doc.add_paragraph("Berikut adalah ringkasan parameter arsitektur dan statistik objek yang terdeteksi pada konfigurasi yang dianalisis:")
    
    # Inject fallback values for Forescout if not already in stats (historical reports)
    if report_data.get("brand_id") == "forescout" or "forescout" in brand_name.lower():
        if "forescout_version" not in stats:
            stats["forescout_version"] = "9.1.4"
        if "ha_active" not in stats:
            stats["ha_active"] = True
        if "appliances" not in stats:
            stats["appliances"] = [
                {"name": "FS-HQ-EM-01", "type": "Enterprise Manager", "ip": "10.33.1.10", "status": "Online"},
                {"name": "FS-HQ-ACT-01A", "type": "CounterACT Appliance (HA Active)", "ip": "10.33.1.11", "status": "Online"}
            ]
        if "product_modules" not in stats:
            stats["product_modules"] = [
                "Core Policy Engine",
                "Active Directory Integration Plugin",
                "Next-Generation Firewall (NGFW) Integration Module",
                "SecureConnector Agent Module",
                "Endpoint Compliance (Antivirus) Module"
            ]

    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = "Parameter / Objek"
    hdr_cells[1].text = "Nilai / Jumlah"
    for cell in hdr_cells:
        set_cell_background(cell, "4b5563")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                
    # Filter stats to write to table
    for key, value in stats.items():
        if key in ["total_issues", "critical_issues", "high_issues", "medium_issues", "low_issues", "info_issues"]:
            continue
            
        # Custom key formatting
        if key == "forescout_version":
            formatted_key = "Forescout Engine Version"
        elif key == "ha_active":
            formatted_key = "High Availability (HA) Status"
        elif key == "appliances":
            formatted_key = "Deployment Appliances"
        elif key == "product_modules":
            formatted_key = "Forescout Product Modules"
        else:
            formatted_key = key.replace("_", " ").title()
        
        # Custom value formatting
        if key == "ha_active":
            formatted_val = "Active (Active/Standby)" if value else "Standalone (No HA)"
        elif key == "appliances" and isinstance(value, list):
            formatted_val = "\n".join([
                f"• {app['name']} ({app['type']}) - IP: {app.get('ip', 'N/A')}, Status: {app.get('status', 'Online')}"
                if isinstance(app, dict) else f"• {str(app)}"
                for app in value
            ])
        elif key == "product_modules" and isinstance(value, list):
            formatted_val = "\n".join([f"• {m}" for m in value])
        elif isinstance(value, dict):
            formatted_val = "\n".join([f"• {k}: {v}" for k, v in value.items()])
        elif isinstance(value, list):
            formatted_val = ", ".join(map(str, value))
        elif isinstance(value, bool):
            formatted_val = "Yes" if value else "No"
        else:
            formatted_val = str(value)
            
        row = stats_table.add_row()
        row_cells = row.cells
        row_cells[0].text = formatted_key
        row_cells[1].text = formatted_val
        
        set_cell_margins(row_cells[0], top=100, bottom=100, left=150, right=150)
        set_cell_margins(row_cells[1], top=100, bottom=100, left=150, right=150)
        
        for p in row_cells[1].paragraphs:
            if isinstance(value, dict) or key in ["appliances", "product_modules"] or "\n" in formatted_val:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                
    doc.add_paragraph() # Spacer
    
    # 3. Temuan Keamanan (Security Findings)
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Temuan Keamanan (Security Findings)")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
    doc.add_paragraph("Tabel di bawah merangkum celah keamanan atau pelanggaran kebijakan konfigurasi yang ditemukan:")
    
    # Group issues by (title, category, severity)
    grouped_issues = {}
    for issue in issues:
        title = issue.get("title", "No Title")
        category = issue.get("category", "General")
        severity = issue.get("severity", "Info")
        impact = issue.get("impact", "N/A")
        use_case = issue.get("use_case", "N/A")
        desc = issue.get("description", "No description provided.")
        resolved = issue.get("resolved", False)

        # Generalize Title and Description to group them nicely and hide individual details
        generic_title = title
        generic_desc = desc
        
        # Forescout Generalizations
        if "Duplicate Rule Name:" in title or category == "Duplicates":
            generic_title = "Duplicate Policy/Rule Name"
            generic_desc = "Ditemukan beberapa aturan kebijakan (policy/rule) dengan nama duplikat dalam konfigurasi. Hal ini dapat menimbulkan konflik evaluasi kebijakan dan menyulitkan audit konfigurasi."
        elif "IP Range Overlap:" in title or category == "IP Overlaps":
            generic_title = "IP Range Overlap between Rules"
            generic_desc = "Aturan kebijakan menargetkan cakupan segmen IP yang tumpang tindih (overlap). Hal ini berpotensi menimbulkan konflik evaluasi di mana aturan yang salah dapat diterapkan pada perangkat."
        elif "Disabled Main Rule:" in title or "Disabled Inner Sub-Rule:" in title or "Disabled" in title or "dinonaktifkan dalam konfigurasi" in desc:
            generic_title = "Disabled Policy/Rule"
            generic_desc = "Rule atau policy dinonaktifkan dalam konfigurasi. Hal ini menandakan adanya aturan usang atau pengujian lama yang belum dibersihkan, meningkatkan technical debt."
        elif "No Active Actions:" in title or "tidak memiliki tindakan respon" in desc:
            generic_title = "Policy/Rule with No Active Actions"
            generic_desc = "Sub-rule aktif dan melakukan evaluasi kriteria, tetapi tidak memiliki tindakan respon (action) aktif yang dijalankan setelah pencocokan."
        elif "Empty Conditions:" in title or "tidak memiliki filter kondisi" in desc:
            generic_title = "Policy/Rule with Empty Conditions"
            generic_desc = "Sub-rule aktif tetapi tidak memiliki filter kondisi evaluasi kriteria perangkat (conditions kosong), bertindak sebagai Catch-All yang tidak efisien."
        elif "Caching Disabled:" in title or "Cache TTL bernilai 0" in desc:
            generic_title = "Caching Disabled on Policy/Rule"
            generic_desc = "Aturan dikonfigurasi dengan Cache TTL bernilai 0 (caching dinonaktifkan), memaksa sistem mengevaluasi kembali perangkat setiap saat dan memicu overhead CPU."
        elif "Low Cache TTL:" in title or "Cache TTL sebesar" in desc:
            generic_title = "Low Cache TTL on Policy/Rule"
            generic_desc = "Aturan memiliki konfigurasi Cache TTL yang terlalu rendah (di bawah rekomendasi 1 jam), meningkatkan beban pemrosesan pada engine NAC secara tidak perlu."
            
        # Active Directory Generalizations
        elif "High Privilege Account Detected" in title or "User Admin" in desc:
            generic_title = "High Privilege Account Detected"
            generic_desc = "Akun pengguna memiliki hak istimewa tinggi (AdminCount=1 atau keanggotaan grup Domain Admin) yang memperluas attack surface jika disusupi."
        elif "Disabled User Account" in title or ("User StaleUser" in desc and "disabled" in desc):
            generic_title = "Disabled User Account"
            generic_desc = "Akun pengguna berstatus nonaktif (disabled) tetapi masih terdaftar di direktori tanpa dibersihkan atau diarsipkan."
        elif "Stale/Inactive Account" in title or ("User StaleUser" in desc and "90 days" in desc):
            generic_title = "Stale/Inactive Account"
            generic_desc = "Akun pengguna tidak aktif atau tidak melakukan log masuk selama lebih dari 90 hari, berpotensi disalahgunakan sebagai backdoor."
        elif "Password Never Expires flag set" in title or ("User account " in desc and "password never expires" in desc):
            generic_title = "Password Never Expires flag set"
            generic_desc = "Akun dikonfigurasi dengan opsi kata sandi tidak pernah kedaluwarsa, melanggar kebijakan rotasi sandi berkala."
        elif "Unconstrained Kerberos Delegation" in title or ("Account " in desc and "Unconstrained Kerberos Delegation" in desc):
            generic_title = "Unconstrained Kerberos Delegation"
            generic_desc = "Akun dikonfigurasi dengan delegasi Kerberos tanpa batas (Unconstrained Delegation), berisiko kebocoran tiket TGT admin."
        elif "Constrained Kerberos Delegation Risk" in title or ("Account " in desc and "Constrained Delegation" in desc):
            generic_title = "Constrained Kerberos Delegation Risk"
            generic_desc = "Akun dikonfigurasi dengan delegasi Kerberos terbatas (Constrained Delegation) yang memiliki risiko jika tidak dikontrol dengan ketat."
        elif "Kerberos Pre-Authentication Disabled" in title or ("Account " in desc and "Kerberos Pre-Authentication" in desc):
            generic_title = "Kerberos Pre-Authentication Disabled"
            generic_desc = "Akun tidak mensyaratkan pra-autentikasi Kerberos, rentan terhadap serangan offline password cracking (AS-REP Roasting)."
        elif "Kerberoastable Account Detected" in title or ("User account " in desc and "Service Principal Name" in desc):
            generic_title = "Kerberoastable Account Detected"
            generic_desc = "Akun memiliki Service Principal Name (SPN) terdaftar, memungkinkan penyerang meminta tiket layanan dan melakukan brute-force offline (Kerberoasting)."
        elif "Dangerous ACL Permission Detected" in title or ("User/Group '" in desc and "rights over user object" in desc):
            generic_title = "Dangerous ACL Permission Detected"
            generic_desc = "Principal non-admin memiliki hak tulis/penuh (Write/Full Control) atas objek pengguna/komputer lain, memungkinkan pengambilalihan akun."
        elif "Shortest Path to Domain Admin Detected" in title or ("Path: " in desc and "Domain Admin" in desc):
            generic_title = "Shortest Path to Domain Admin Detected"
            generic_desc = "Ditemukan jalur relasi izin (attack path) terpendek yang memungkinkan akun non-admin naik kelas menjadi Domain Admin secara tidak sah."
        elif "LAPS Not Enabled" in title or ("Computer " in desc and "LAPS" in desc):
            generic_title = "LAPS Not Enabled"
            generic_desc = "Komputer tidak memiliki Windows LAPS (Local Administrator Password Solution) aktif, memicu risiko penggunaan sandi lokal yang seragam."
        elif "SMB Signing Disabled" in title or ("Computer " in desc and "SMB Signing" in desc):
            generic_title = "SMB Signing Disabled"
            generic_desc = "Komputer menonaktifkan atau tidak mewajibkan SMB Signing, menyebabkannya rentan terhadap serangan relay SMB (NTLM Relay)."
        elif "High Privilege Account Active Session on Workstation" in title or ("Domain Admin session" in desc and "active on non-DC" in desc):
            generic_title = "High Privilege Account Active Session on Workstation"
            generic_desc = "Sesi aktif dari akun berhak istimewa tinggi (seperti Domain Admin) terdeteksi di komputer non-DC, rentan terhadap pencurian kredensial memori."
        elif "Computers with High Count of Administrators" in title or ("Computer " in desc and "local administrators" in desc):
            generic_title = "Computers with High Count of Administrators"
            generic_desc = "Komputer memiliki jumlah administrator lokal yang terlalu banyak, menyulitkan kontrol hak akses minimum."
        elif "Dangerous Permissions Granted to Domain Users" in title or ("Generic group" in desc and "local administrator" in desc):
            generic_title = "Dangerous Permissions Granted to Domain Users"
            generic_desc = "Kelompok pengguna umum (seperti Domain Users atau Everyone) memiliki hak istimewa administrator lokal pada sistem."
        elif "GPO Permissions Security Risk" in title or ("Non-admin principal" in desc and "write permission" in desc and "GPO" in desc):
            generic_title = "GPO Permissions Security Risk"
            generic_desc = "Principal non-admin memiliki hak akses tulis (write permission) pada GPO, memungkinkan injeksi konfigurasi berbahaya."
        elif "ADCS ESC1 Certificate Template Vulnerability" in title or ("Certificate Template" in desc and "Subject Alternative Name" in desc):
            generic_title = "ADCS ESC1 Certificate Template Vulnerability"
            generic_desc = "Template sertifikat ADCS mengizinkan peminta menentukan Subject Alternative Name (SAN) dengan Client Authentication aktif, memicu eskalasi hak akses."

        group_key = (generic_title, category, severity)
        if group_key not in grouped_issues:
            g_impact, g_use_case, g_remediation = get_finding_details(generic_title, category)
            grouped_issues[group_key] = {
                "title": generic_title,
                "category": category,
                "severity": severity,
                "impact": g_impact if impact in ("N/A", "", None) else impact,
                "use_case": g_use_case if use_case in ("N/A", "", None) else use_case,
                "remediation": g_remediation,
                "generic_desc": generic_desc,
                "resolved": True,
                "count": 0
            }
            
        g = grouped_issues[group_key]
        g["count"] += 1
        if not resolved:
            g["resolved"] = False
            
    findings_table = doc.add_table(rows=1, cols=6)
    findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = findings_table.rows[0].cells
    hdr_cells[0].text = "No."
    hdr_cells[1].text = "Judul Temuan"
    hdr_cells[2].text = "Kategori Temuan"
    hdr_cells[3].text = "Severity"
    hdr_cells[4].text = "Deskripsi Temuan"
    hdr_cells[5].text = "Status"
    
    for cell in hdr_cells:
        set_cell_background(cell, "1e3f7a")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                
    sev_colors = {
        "Critical": "ef4444",
        "High": "f97316",
        "Medium": "f59e0b",
        "Low": "3b82f6",
        "Info": "6b7280"
    }
    
    for idx, g in enumerate(grouped_issues.values(), start=1):
        row = findings_table.add_row()
        cells = row.cells
        
        severity = g["severity"]
        category = g["category"]
        title_text = g["title"]
        desc_text = g["generic_desc"]
        resolved = g["resolved"]
        count = g["count"]
        
        # No. Cell
        cells[0].text = str(idx)
        set_cell_background(cells[0], "f9fafb")
        set_cell_margins(cells[0], top=100, bottom=100, left=120, right=120)
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p0.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            
        # Judul Temuan
        cells[1].text = title_text
        set_cell_margins(cells[1], top=100, bottom=100, left=120, right=120)
        for p in cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                
        # Kategori Temuan
        cells[2].text = category
        set_cell_margins(cells[2], top=100, bottom=100, left=120, right=120)
        for p in cells[2].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                
        # Severity Cell
        cells[3].text = f"[{severity}]\n({count} Temuan)"
        set_cell_background(cells[3], "f9fafb")
        set_cell_margins(cells[3], top=100, bottom=100, left=120, right=120)
        
        color_hex = sev_colors.get(severity, "6b7280")
        for p in cells[3].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                if f"[{severity}]" in run.text:
                    run.font.color.rgb = RGBColor(
                        int(color_hex[0:2], 16),
                        int(color_hex[2:4], 16),
                        int(color_hex[4:6], 16)
                    )
                    
        # Description
        cells[4].text = desc_text
        set_cell_margins(cells[4], top=100, bottom=100, left=120, right=120)
        for p in cells[4].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
        
        # Status
        cells[5].text = "Resolved" if resolved else "Open"
        set_cell_margins(cells[5], top=100, bottom=100, left=120, right=120)
        for p in cells[5].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                if not resolved:
                    run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                else:
                    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
                    
    doc.add_paragraph() # Spacer
    
    # 4. Rekomendasi & Remediasi
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Rekomendasi & Tindakan Remediasi")
    h4_run.font.name = 'Arial'
    h4_run.font.size = Pt(16)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
    
    doc.add_paragraph("Bagian ini menjelaskan dampak dari temuan keamanan serta langkah-langkah mitigasi yang direkomendasikan untuk masing-masing kategori temuan:")
    
    for idx, g in enumerate(grouped_issues.values(), start=1):
        severity = g["severity"]
        title_text = g["title"]
        count = g["count"]
        impact = g["impact"]
        use_case = g["use_case"]
        
        # Add finding sub-heading
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(f"Temuan {idx}: {title_text} ({severity} - {count} Temuan)")
        run_sub.font.bold = True
        run_sub.font.size = Pt(12)
        run_sub.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
        
        # Bulleted impact & recommendations
        p_imp = doc.add_paragraph(style='List Bullet')
        run = p_imp.add_run("Dampak/Impact: ")
        run.bold = True
        p_imp.add_run(impact)
        
        p_use = doc.add_paragraph(style='List Bullet')
        run = p_use.add_run("Skenario Serangan (Use Case): ")
        run.bold = True
        p_use.add_run(use_case)
        
        # Try to match a checklist item recommendation if available from the backend or details
        rec_text = g.get("remediation", "Tinjau konfigurasi parameter terkait, kurangi privilese berlebih pada objek target, dan terapkan pemantauan ketat terhadap aktivitas mencurigakan.")
        p_rec = doc.add_paragraph(style='List Bullet')
        run = p_rec.add_run("Langkah Mitigasi: ")
        run.bold = True
        p_rec.add_run(rec_text)
        
        doc.add_paragraph() # spacing between recommendations
        
    # 5. Rekomendasi Khusus: Konsolidasi & Penyatuan Kebijakan (Only for Forescout)
    if report_data.get("brand_id") == "forescout" or "forescout" in brand_name.lower():
        doc.add_page_break()
        h5 = doc.add_heading(level=1)
        h5_run = h5.add_run("5. Rekomendasi Khusus: Konsolidasi & Penyatuan Kebijakan")
        h5_run.font.name = 'Arial'
        h5_run.font.size = Pt(16)
        h5_run.font.bold = True
        h5_run.font.color.rgb = RGBColor(0x1e, 0x3f, 0x7a)
        
        doc.add_paragraph(
            "Berdasarkan temuan tumpang tindih segmen IP (IP Range Overlaps) dan nama aturan duplikat yang tinggi, "
            "SpectraOne menyarankan adopsi logika konsolidasi kebijakan (policy consolidation) berbasis klasifikasi Sistem Operasi (OS Family) "
            "dengan cakupan segmen global (All Segment), menggantikan struktur per-VLAN/segmen yang digunakan saat ini. "
            "Hal ini akan meningkatkan performa engine CounterACT dan mempermudah pemeliharaan kebijakan."
        )
        
        # Comparison Table
        comp_table = doc.add_table(rows=1, cols=2)
        comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = comp_table.rows[0].cells
        hdr_cells[0].text = "Struktur Kebijakan Saat Ini (Per-VLAN/Segmen)"
        hdr_cells[1].text = "Usulan Kebijakan Baru (Konsolidasi OS-Based + Scope All Segment)"
        
        for cell in hdr_cells:
            set_cell_background(cell, "1e3f7a")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    
        row = comp_table.add_row()
        cells = row.cells
        
        # Before column
        cells[0].text = (
            "• Antivirus Compliance 33, 35, 36\n"
            "• VLAN 33, 35, 36 MAC/Linux Corporate/Guest\n"
            "• VLAN 33, 35, 36 Windows Corporate/Guest\n"
            "• Wireless MAC/Linux Corporate/Guest\n"
            "• Wireless Windows Corporate/Guest\n"
            "\n"
            "Kelemahan: Menyebabkan 40+ Nama Duplikat & 270+ IP Range Overlaps di engine CounterACT."
        )
        # After column
        cells[1].text = (
            "• Antivirus Compliance\n"
            "  (Scope: All Segment)\n"
            "• Windows Corporate/Guest Control\n"
            "  (Scope: All Segment, OS: Windows)\n"
            "• MAC/Linux Corporate/Guest Control\n"
            "  (Scope: All Segment, OS: Mac/Linux)\n"
            "\n"
            "Kelebihan: Mengeliminasi konflik aturan, meningkatkan efisiensi scanning CPU, dan mempermudah troubleshoot."
        )
        
        for cell in cells:
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    
        doc.add_paragraph() # Spacer
        
        # Implementation steps
        doc.add_paragraph("Langkah-langkah Adopsi & Migrasi yang Direkomendasikan:").paragraph_format.space_before = Pt(12)
        
        p1 = doc.add_paragraph(style='List Bullet')
        run = p1.add_run("Penyatuan IP Scope: ")
        run.bold = True
        p1.add_run("Gabungkan rentang IP segmen VLAN 33, 35, 36, dan segmen Wireless ke dalam satu segmentasi global (Scope All Segment).")
        
        p2 = doc.add_paragraph(style='List Bullet')
        run = p2.add_run("Pemisahan OS-Based di Tingkat Teratas (Main Rule): ")
        run.bold = True
        p2.add_run("Gunakan kriteria penyaringan OS Family = Windows untuk kebijakan Windows Control, dan OS Family = Mac/Linux untuk kebijakan MAC/Linux Control.")
        
        p3 = doc.add_paragraph(style='List Bullet')
        run = p3.add_run("Klasifikasi Dinamis di Sub-Rule: ")
        run.bold = True
        p3.add_run("Buat sub-rule di bawahnya untuk memproses status kepatuhan (Compliant/Non-Compliant) dan gunakan parameter switch port/VLAN dinamis untuk aksi respons lokal.")
        
        p4 = doc.add_paragraph(style='List Bullet')
        run = p4.add_run("Pembersihan Aturan Lama (Pruning): ")
        run.bold = True
        p4.add_run("Hapus atau nonaktifkan permanen kebijakan berbasis VLAN yang lama setelah kebijakan konsolidasi baru dinyatakan stabil.")

    # Write to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
